import os
import PyPDF2
import pdfplumber
from docx import Document
import streamlit as st
from typing import Optional

class DocumentProcessor:
    """Handles text extraction from various document formats."""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.txt', '.docx']
    
    def extract_text(self, file_path: str, original_filename: str) -> str:
        """
        Extract text from uploaded document based on file extension.
        
        Args:
            file_path: Path to the temporary file
            original_filename: Original filename to determine file type
            
        Returns:
            Extracted text as string
            
        Raises:
            Exception: If file format is unsupported or extraction fails
        """
        try:
            file_extension = os.path.splitext(original_filename)[1].lower()
            
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
            elif file_extension == '.docx':
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            raise Exception(f"Failed to extract text from {original_filename}: {str(e)}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using pdfplumber with PyPDF2 fallback."""
        text = ""
        
        try:
            # Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # If pdfplumber didn't extract much text, try PyPDF2
            if len(text.strip()) < 100:
                text = ""
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                            
        except Exception as e:
            # Fallback to PyPDF2 if pdfplumber fails
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception as fallback_error:
                raise Exception(f"Failed to extract PDF text with both methods: {str(fallback_error)}")
        
        if not text.strip():
            raise Exception("No readable text found in PDF. The document may be scanned or image-based.")
        
        return self._clean_text(text)
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    return self._clean_text(text)
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Unable to decode text file with supported encodings")
            
        except Exception as e:
            raise Exception(f"Failed to read text file: {str(e)}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise Exception("No readable text found in DOCX document")
            
            return self._clean_text(text)
            
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n')]
        lines = [line for line in lines if line]  # Remove empty lines
        
        # Join lines with single spaces, but preserve paragraph breaks
        cleaned_text = ""
        for i, line in enumerate(lines):
            cleaned_text += line
            # Add paragraph break for significant gaps or if line ends with punctuation
            if i < len(lines) - 1:
                if (line.endswith('.') or line.endswith(':') or 
                    line.endswith('!') or line.endswith('?')):
                    cleaned_text += "\n\n"
                else:
                    cleaned_text += " "
        
        return cleaned_text.strip()
    
    def get_document_stats(self, text: str) -> dict:
        """Get basic statistics about the extracted document."""
        if not text:
            return {"words": 0, "characters": 0, "paragraphs": 0}
        
        words = len(text.split())
        characters = len(text)
        paragraphs = len([p for p in text.split('\n\n') if p.strip()])
        
        return {
            "words": words,
            "characters": characters,
            "paragraphs": paragraphs
        }
